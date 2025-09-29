import base64
import os
import secrets
import time
import unicodedata
from typing import Any, Dict, List, Optional, Tuple

import boto3
import botocore
import docx
import gradio as gr
import pandas as pd
import polars as pl
from botocore.client import BaseClient
from faker import Faker
from gradio import Progress
from openpyxl import Workbook
from presidio_analyzer import (
    AnalyzerEngine,
    BatchAnalyzerEngine,
    DictAnalyzerResult,
    RecognizerResult,
)
from presidio_anonymizer import AnonymizerEngine, BatchAnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

from tools.config import (
    AWS_ACCESS_KEY,
    AWS_REGION,
    AWS_SECRET_KEY,
    CUSTOM_ENTITIES,
    DEFAULT_LANGUAGE,
    DO_INITIAL_TABULAR_DATA_CLEAN,
    MAX_SIMULTANEOUS_FILES,
    MAX_TABLE_COLUMNS,
    MAX_TABLE_ROWS,
    OUTPUT_FOLDER,
    PRIORITISE_SSO_OVER_AWS_ENV_ACCESS_KEYS,
    RUN_AWS_FUNCTIONS,
    aws_comprehend_language_choices,
)
from tools.helper_functions import (
    detect_file_type,
    get_file_name_without_type,
    read_file,
)
from tools.load_spacy_model_custom_recognisers import (
    CustomWordFuzzyRecognizer,
    create_nlp_analyser,
    custom_word_list_recogniser,
    load_spacy_model,
    nlp_analyser,
    score_threshold,
)

# Use custom version of analyze_dict to be able to track progress
from tools.presidio_analyzer_custom import analyze_dict
from tools.secure_path_utils import secure_join

custom_entities = CUSTOM_ENTITIES

fake = Faker("en_UK")


def fake_first_name(x):
    return fake.first_name()


# #### Some of my cleaning functions
url_pattern = r"http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+|(?:www\.)[a-zA-Z0-9._-]+\.[a-zA-Z]{2,}"
html_pattern_regex = r"<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});|\xa0|&nbsp;"
html_start_pattern_end_dots_regex = r"<(.*?)\.\."
non_ascii_pattern = r"[^\x00-\x7F]+"
and_sign_regex = r"&"
multiple_spaces_regex = r"\s{2,}"
multiple_new_lines_regex = r"(\r\n|\n)+"
multiple_punctuation_regex = r"(\p{P})\p{P}+"


def initial_clean(texts: pd.Series) -> pd.Series:
    """
    This function cleans the text by removing URLs, HTML tags, and non-ASCII characters.
    """
    for text in texts:
        if not text or pd.isnull(text):
            text = ""

        # Normalize unicode characters to decompose any special forms
        normalized_text = unicodedata.normalize("NFKC", text)

        # Replace smart quotes and special punctuation with standard ASCII equivalents
        replacements = {
            "‘": "'",
            "’": "'",
            "“": '"',
            "”": '"',
            "–": "-",
            "—": "-",
            "…": "...",
            "•": "*",
        }

        # Perform replacements
        for old_char, new_char in replacements.items():
            normalised_text = normalized_text.replace(old_char, new_char)

        text = normalised_text

    # Convert to polars Series
    texts = pl.Series(texts).str.strip_chars()

    # Define a list of patterns and their replacements
    patterns = [
        (multiple_new_lines_regex, "  "),
        (r"\r", ""),
        (url_pattern, " "),
        (html_pattern_regex, " "),
        (html_start_pattern_end_dots_regex, " "),
        (non_ascii_pattern, " "),
        (multiple_spaces_regex, " "),
        (multiple_punctuation_regex, "${1}"),
        (and_sign_regex, "and"),
    ]

    # Apply each regex replacement
    for pattern, replacement in patterns:
        texts = texts.str.replace_all(pattern, replacement)

    # Convert the series back to a list
    texts = texts.to_list()

    return texts


def process_recognizer_result(
    result: RecognizerResult,
    recognizer_result: RecognizerResult,
    data_row: int,
    dictionary_key: int,
    df_dict: Dict[str, List[Any]],
    keys_to_keep: List[str],
) -> Tuple[List[str], List[Dict[str, Any]]]:
    output = list()
    output_dicts = list()

    if hasattr(result, "value"):
        text = result.value[data_row]
    else:
        text = ""

    if isinstance(recognizer_result, list):
        for sub_result in recognizer_result:
            if isinstance(text, str):
                found_text = text[sub_result.start : sub_result.end]
            else:
                found_text = ""
            analysis_explanation = {
                key: sub_result.__dict__[key] for key in keys_to_keep
            }
            analysis_explanation.update(
                {
                    "data_row": str(data_row),
                    "column": list(df_dict.keys())[dictionary_key],
                    "entity": found_text,
                }
            )
            output.append(str(analysis_explanation))
            output_dicts.append(analysis_explanation)

    return output, output_dicts


# Writing decision making process to file
def generate_log(
    analyzer_results: List[DictAnalyzerResult], df_dict: Dict[str, List[Any]]
) -> Tuple[str, pd.DataFrame]:
    """
    Generate a detailed output of the decision process for entity recognition.

    This function takes the results from the analyzer and the original data dictionary,
    and produces a string output detailing the decision process for each recognized entity.
    It includes information such as entity type, position, confidence score, and the context
    in which the entity was found.

    Args:
        analyzer_results (List[DictAnalyzerResult]): The results from the entity analyzer.
        df_dict (Dict[str, List[Any]]): The original data in dictionary format.

    Returns:
        Tuple[str, pd.DataFrame]: A tuple containing the string output and DataFrame with all columns.
    """
    decision_process_output = list()
    decision_process_output_dicts = list()  # New list to store dictionaries
    keys_to_keep = ["entity_type", "start", "end"]

    # Run through each column to analyse for PII
    for i, result in enumerate(analyzer_results):

        # If a single result
        if isinstance(result, RecognizerResult):
            output, output_dicts = process_recognizer_result(
                result, result, 0, i, df_dict, keys_to_keep
            )
            decision_process_output.extend(output)
            decision_process_output_dicts.extend(output_dicts)

        # If a list of results
        elif isinstance(result, list) or isinstance(result, DictAnalyzerResult):
            for x, recognizer_result in enumerate(result.recognizer_results):
                output, output_dicts = process_recognizer_result(
                    result, recognizer_result, x, i, df_dict, keys_to_keep
                )
                decision_process_output.extend(output)
                decision_process_output_dicts.extend(output_dicts)

        else:
            try:
                output, output_dicts = process_recognizer_result(
                    result, result, 0, i, df_dict, keys_to_keep
                )
                decision_process_output.extend(output)
                decision_process_output_dicts.extend(output_dicts)
            except Exception as e:
                print(e)

    decision_process_output_str = "\n".join(decision_process_output)
    decision_process_output_df = pd.DataFrame(decision_process_output_dicts)

    return decision_process_output_str, decision_process_output_df


def anon_consistent_names(df: pd.DataFrame) -> pd.DataFrame:
    # ## Pick out common names and replace them with the same person value
    df_dict = df.to_dict(orient="list")

    # analyzer = AnalyzerEngine()
    batch_analyzer = BatchAnalyzerEngine(analyzer_engine=nlp_analyser)

    analyzer_results = batch_analyzer.analyze_dict(df_dict, language=DEFAULT_LANGUAGE)
    analyzer_results = list(analyzer_results)

    text = analyzer_results[3].value

    recognizer_result = str(analyzer_results[3].recognizer_results)

    data_str = recognizer_result  # abbreviated for brevity

    # Adjusting the parse_dict function to handle trailing ']'
    # Splitting the main data string into individual list strings
    list_strs = data_str[1:-1].split("], [")

    def parse_dict(s):
        s = s.strip("[]")  # Removing any surrounding brackets
        items = s.split(", ")
        d = {}
        for item in items:
            key, value = item.split(": ")
            if key == "score":
                d[key] = float(value)
            elif key in ["start", "end"]:
                d[key] = int(value)
            else:
                d[key] = value
        return d

    # Re-running the improved processing code

    result = list()

    for lst_str in list_strs:
        # Splitting each list string into individual dictionary strings
        dict_strs = lst_str.split(", type: ")
        dict_strs = [dict_strs[0]] + [
            "type: " + s for s in dict_strs[1:]
        ]  # Prepending "type: " back to the split strings

        # Parsing each dictionary string
        dicts = [parse_dict(d) for d in dict_strs]
        result.append(dicts)

    names = list()

    for idx, paragraph in enumerate(text):
        paragraph_texts = list()
        for dictionary in result[idx]:
            if dictionary["type"] == "PERSON":
                paragraph_texts.append(
                    paragraph[dictionary["start"] : dictionary["end"]]
                )
        names.append(paragraph_texts)

    # Flatten the list of lists and extract unique names
    unique_names = list(set(name for sublist in names for name in sublist))

    fake_names = pd.Series(unique_names).apply(fake_first_name)

    mapping_df = pd.DataFrame(
        data={"Unique names": unique_names, "Fake names": fake_names}
    )

    # Convert mapping dataframe to dictionary, adding word boundaries for full-word match
    name_map = {
        r"\b" + k + r"\b": v
        for k, v in zip(mapping_df["Unique names"], mapping_df["Fake names"])
    }

    name_map

    scrubbed_df_consistent_names = df.replace(name_map, regex=True)

    scrubbed_df_consistent_names

    return scrubbed_df_consistent_names


def handle_docx_anonymisation(
    file_path: str,
    output_folder: str,
    anon_strategy: str,
    chosen_redact_entities: List[str],
    in_allow_list: List[str],
    in_deny_list: List[str],
    max_fuzzy_spelling_mistakes_num: int,
    pii_identification_method: str,
    chosen_redact_comprehend_entities: List[str],
    comprehend_query_number: int,
    comprehend_client: BaseClient,
    language: Optional[str] = DEFAULT_LANGUAGE,
    out_file_paths: List[str] = list(),
    nlp_analyser: AnalyzerEngine = nlp_analyser,
):
    """
    Anonymises a .docx file by extracting text, processing it, and re-inserting it.

    Returns:
        A tuple containing the output file path and the log file path.
    """

    # 1. Load the document and extract text elements
    doc = docx.Document(file_path)
    text_elements = (
        list()
    )  # This will store the actual docx objects (paragraphs, cells)
    original_texts = list()  # This will store the text from those objects

    paragraph_count = len(doc.paragraphs)

    if paragraph_count > MAX_TABLE_ROWS:
        out_message = f"Number of paragraphs in document is greater than {MAX_TABLE_ROWS}. Please submit a smaller document."
        print(out_message)
        raise Exception(out_message)

    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            text_elements.append(para)
            original_texts.append(para.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Only process non-empty cells
                    text_elements.append(cell)
                    original_texts.append(cell.text)

    # If there's no text to process, return early
    if not original_texts:
        print(f"No text found in {file_path}. Skipping.")
        return None, None, 0

    # 2. Convert to a DataFrame for the existing anonymisation script
    df_to_anonymise = pd.DataFrame({"text_to_redact": original_texts})

    # 3. Call the core anonymisation script
    (
        anonymised_df,
        _,
        decision_log,
        comprehend_query_number,
        decision_process_output_df,
    ) = anonymise_script(
        df=df_to_anonymise,
        anon_strategy=anon_strategy,
        language=language,
        chosen_redact_entities=chosen_redact_entities,
        in_allow_list=in_allow_list,
        in_deny_list=in_deny_list,
        max_fuzzy_spelling_mistakes_num=max_fuzzy_spelling_mistakes_num,
        pii_identification_method=pii_identification_method,
        chosen_redact_comprehend_entities=chosen_redact_comprehend_entities,
        comprehend_query_number=comprehend_query_number,
        comprehend_client=comprehend_client,
        nlp_analyser=nlp_analyser,
    )

    anonymised_texts = anonymised_df["text_to_redact"].tolist()

    # 4. Re-insert the anonymised text back into the document objects
    for element, new_text in zip(text_elements, anonymised_texts):
        if isinstance(element, docx.text.paragraph.Paragraph):
            # Clear existing content (runs) and add the new text in a single new run
            element.clear()
            element.add_run(new_text)
        elif isinstance(element, docx.table._Cell):
            # For cells, setting .text works similarly
            element.text = new_text

    # 5. Save the redacted document and the log file
    base_name = os.path.basename(file_path)
    file_name_without_ext = os.path.splitext(base_name)[0]

    output_docx_path = secure_join(
        output_folder, f"{file_name_without_ext}_redacted.docx"
    )

    out_file_paths.append(output_docx_path)

    output_xlsx_path = secure_join(
        output_folder, f"{file_name_without_ext}_redacted.csv"
    )

    anonymised_df.to_csv(output_xlsx_path, encoding="utf-8-sig", index=None)
    doc.save(output_docx_path)

    out_file_paths.append(output_xlsx_path)

    # Reconstruct log_file_path for return value
    log_file_path = secure_join(
        output_folder, f"{file_name_without_ext}_redacted_log.csv"
    )

    decision_process_output_df.to_csv(log_file_path, index=None, encoding="utf-8-sig")

    out_file_paths.append(log_file_path)

    return out_file_paths, comprehend_query_number


def anonymise_files_with_open_text(
    file_paths: List[str],
    in_text: str,
    anon_strategy: str,
    chosen_cols: List[str],
    chosen_redact_entities: List[str],
    in_allow_list: List[str] = None,
    latest_file_completed: int = 0,
    out_message: list = list(),
    out_file_paths: list = list(),
    log_files_output_paths: list = list(),
    in_excel_sheets: list = list(),
    first_loop_state: bool = False,
    output_folder: str = OUTPUT_FOLDER,
    in_deny_list: list[str] = list(),
    max_fuzzy_spelling_mistakes_num: int = 0,
    pii_identification_method: str = "Local",
    chosen_redact_comprehend_entities: List[str] = list(),
    comprehend_query_number: int = 0,
    aws_access_key_textbox: str = "",
    aws_secret_key_textbox: str = "",
    actual_time_taken_number: float = 0,
    do_initial_clean: bool = DO_INITIAL_TABULAR_DATA_CLEAN,
    language: Optional[str] = None,
    progress: Progress = Progress(track_tqdm=True),
):
    """
    This function anonymises data files based on the provided parameters.

    Parameters:
    - file_paths (List[str]): A list of file paths to anonymise: '.xlsx', '.xls', '.csv', '.parquet', or '.docx'.
    - in_text (str): The text to anonymise if file_paths is 'open_text'.
    - anon_strategy (str): The anonymisation strategy to use.
    - chosen_cols (List[str]): A list of column names to anonymise.
    - language (str): The language of the text to anonymise.
    - chosen_redact_entities (List[str]): A list of entities to redact.
    - in_allow_list (List[str], optional): A list of allowed values. Defaults to None.
    - latest_file_completed (int, optional): The index of the last file completed. Defaults to 0.
    - out_message (list, optional): A list to store output messages. Defaults to an empty list.
    - out_file_paths (list, optional): A list to store output file paths. Defaults to an empty list.
    - log_files_output_paths (list, optional): A list to store log file paths. Defaults to an empty list.
    - in_excel_sheets (list, optional): A list of Excel sheet names. Defaults to an empty list.
    - first_loop_state (bool, optional): Indicates if this is the first loop iteration. Defaults to False.
    - output_folder (str, optional): The output folder path. Defaults to the global output_folder variable.
    - in_deny_list (list[str], optional): A list of specific terms to redact.
    - max_fuzzy_spelling_mistakes_num (int, optional): The maximum number of spelling mistakes allowed in a searched phrase for fuzzy matching. Can range from 0-9.
    - pii_identification_method (str, optional): The method to redact personal information. Either 'Local' (spacy model), or 'AWS Comprehend' (AWS Comprehend API).
    - chosen_redact_comprehend_entities (List[str]): A list of entity types to redact from files, chosen from the official list from AWS Comprehend service.
    - comprehend_query_number (int, optional): A counter tracking the number of queries to AWS Comprehend.
    - aws_access_key_textbox (str, optional): AWS access key for account with Textract and Comprehend permissions.
    - aws_secret_key_textbox (str, optional): AWS secret key for account with Textract and Comprehend permissions.
    - actual_time_taken_number (float, optional): Time taken to do the redaction.
    - language (str, optional): The language of the text to anonymise.
    - progress (Progress, optional): A Progress object to track progress. Defaults to a Progress object with track_tqdm=True.
    - do_initial_clean (bool, optional): Whether to perform an initial cleaning of the text. Defaults to True.
    """

    tic = time.perf_counter()
    comprehend_client = ""

    # If output folder doesn't end with a forward slash, add one
    if not output_folder.endswith("/"):
        output_folder = output_folder + "/"

    # Use provided language or default
    language = language or DEFAULT_LANGUAGE

    if pii_identification_method == "AWS Comprehend":
        if language not in aws_comprehend_language_choices:
            out_message = f"Please note that this language is not supported by AWS Comprehend: {language}"
            raise Warning(out_message)

    # If this is the first time around, set variables to 0/blank
    if first_loop_state is True:
        latest_file_completed = 0
        out_message = list()
        out_file_paths = list()

    # Load file
    # If out message or out_file_paths are blank, change to a list so it can be appended to
    if isinstance(out_message, str):
        out_message = [out_message]

    if isinstance(log_files_output_paths, str):
        log_files_output_paths = list()

    if not out_file_paths:
        out_file_paths = list()

    if isinstance(in_allow_list, list):
        if in_allow_list:
            in_allow_list_flat = in_allow_list
        else:
            in_allow_list_flat = list()
    elif isinstance(in_allow_list, pd.DataFrame):
        if not in_allow_list.empty:
            in_allow_list_flat = list(in_allow_list.iloc[:, 0].unique())
        else:
            in_allow_list_flat = list()
    else:
        in_allow_list_flat = list()

    anon_df = pd.DataFrame()

    # Try to connect to AWS services directly only if RUN_AWS_FUNCTIONS environmental variable is 1, otherwise an environment variable or direct textbox input is needed.
    if pii_identification_method == "AWS Comprehend":
        print("Trying to connect to AWS Comprehend service")
        if RUN_AWS_FUNCTIONS == "1" and PRIORITISE_SSO_OVER_AWS_ENV_ACCESS_KEYS == "1":
            print("Connecting to Comprehend via existing SSO connection")
            comprehend_client = boto3.client("comprehend", region_name=AWS_REGION)
        elif aws_access_key_textbox and aws_secret_key_textbox:
            print(
                "Connecting to Comprehend using AWS access key and secret keys from textboxes."
            )
            comprehend_client = boto3.client(
                "comprehend",
                aws_access_key_id=aws_access_key_textbox,
                aws_secret_access_key=aws_secret_key_textbox,
            )
        elif RUN_AWS_FUNCTIONS == "1":
            print("Connecting to Comprehend via existing SSO connection")
            comprehend_client = boto3.client("comprehend")
        elif AWS_ACCESS_KEY and AWS_SECRET_KEY:
            print("Getting Comprehend credentials from environment variables")
            comprehend_client = boto3.client(
                "comprehend",
                aws_access_key_id=AWS_ACCESS_KEY,
                aws_secret_access_key=AWS_SECRET_KEY,
            )
        else:
            comprehend_client = ""
            out_message = "Cannot connect to AWS Comprehend service. Please provide access keys under Textract settings on the Redaction settings tab, or choose another PII identification method."
            raise (out_message)

    # Check if files and text exist
    if not file_paths:
        if in_text:
            file_paths = ["open_text"]
        else:
            out_message = "Please enter text or a file to redact."
            raise Exception(out_message)

    if not isinstance(file_paths, list):
        file_paths = [file_paths]

    if len(file_paths) > MAX_SIMULTANEOUS_FILES:
        out_message = f"Number of files to anonymise is greater than {MAX_SIMULTANEOUS_FILES}. Please submit a smaller number of files."
        print(out_message)
        raise Exception(out_message)

    # If we have already redacted the last file, return the input out_message and file list to the relevant components
    if latest_file_completed >= len(file_paths):
        print("Last file reached")  # , returning files:", str(latest_file_completed))
        # Set to a very high number so as not to mess with subsequent file processing by the user
        # latest_file_completed = 99
        final_out_message = "\n".join(out_message)

        gr.Info(final_out_message)

        return (
            final_out_message,
            out_file_paths,
            out_file_paths,
            latest_file_completed,
            log_files_output_paths,
            log_files_output_paths,
            actual_time_taken_number,
            comprehend_query_number,
        )

    file_path_loop = [file_paths[int(latest_file_completed)]]

    for anon_file in progress.tqdm(
        file_path_loop, desc="Anonymising files", unit="files"
    ):

        # Get a string file path
        if isinstance(anon_file, str):
            file_path = anon_file
        else:
            file_path = anon_file

        if anon_file == "open_text":
            anon_df = pd.DataFrame(data={"text": [in_text]})
            chosen_cols = ["text"]
            out_file_part = anon_file
            sheet_name = ""
            file_type = ""

            (
                out_file_paths,
                out_message,
                key_string,
                log_files_output_paths,
                comprehend_query_number,
            ) = tabular_anonymise_wrapper_func(
                file_path,
                anon_df,
                chosen_cols,
                out_file_paths,
                out_file_part,
                out_message,
                sheet_name,
                anon_strategy,
                language,
                chosen_redact_entities,
                in_allow_list,
                file_type,
                "",
                log_files_output_paths,
                in_deny_list,
                max_fuzzy_spelling_mistakes_num,
                pii_identification_method,
                chosen_redact_comprehend_entities,
                comprehend_query_number,
                comprehend_client,
                output_folder=OUTPUT_FOLDER,
                do_initial_clean=do_initial_clean,
            )
        else:
            # If file is an xlsx, we are going to run through all the Excel sheets to anonymise them separately.
            file_type = detect_file_type(file_path)
            print("File type is:", file_type)

            out_file_part = get_file_name_without_type(file_path)

            if file_type == "docx":
                out_file_paths, comprehend_query_number = handle_docx_anonymisation(
                    file_path=file_path,
                    output_folder=output_folder,
                    anon_strategy=anon_strategy,
                    chosen_redact_entities=chosen_redact_entities,
                    in_allow_list=in_allow_list_flat,
                    in_deny_list=in_deny_list,
                    max_fuzzy_spelling_mistakes_num=max_fuzzy_spelling_mistakes_num,
                    pii_identification_method=pii_identification_method,
                    chosen_redact_comprehend_entities=chosen_redact_comprehend_entities,
                    comprehend_query_number=comprehend_query_number,
                    comprehend_client=comprehend_client,
                    language=language,
                    out_file_paths=out_file_paths,
                )

            elif file_type == "xlsx":
                print("Running through all xlsx sheets")
                # anon_xlsx = pd.ExcelFile(anon_file)
                if not in_excel_sheets:
                    out_message.append(
                        "No Excel sheets selected. Please select at least one to anonymise."
                    )
                    continue

                # Create xlsx file:
                anon_xlsx = pd.ExcelFile(file_path)
                anon_xlsx_export_file_name = (
                    output_folder + out_file_part + "_redacted.xlsx"
                )

                # Iterate through the sheet names
                for sheet_name in progress.tqdm(
                    in_excel_sheets, desc="Anonymising sheets", unit="sheets"
                ):
                    # Read each sheet into a DataFrame
                    if sheet_name not in anon_xlsx.sheet_names:
                        continue

                    anon_df = pd.read_excel(file_path, sheet_name=sheet_name)

                    (
                        out_file_paths,
                        out_message,
                        key_string,
                        log_files_output_paths,
                        comprehend_query_number,
                    ) = tabular_anonymise_wrapper_func(
                        anon_file,
                        anon_df,
                        chosen_cols,
                        out_file_paths,
                        out_file_part,
                        out_message,
                        sheet_name,
                        anon_strategy,
                        language,
                        chosen_redact_entities,
                        in_allow_list,
                        file_type,
                        anon_xlsx_export_file_name,
                        log_files_output_paths,
                        in_deny_list,
                        max_fuzzy_spelling_mistakes_num,
                        pii_identification_method,
                        language,
                        chosen_redact_comprehend_entities,
                        comprehend_query_number,
                        comprehend_client,
                        output_folder=output_folder,
                        do_initial_clean=do_initial_clean,
                    )

            else:
                sheet_name = ""
                anon_df = read_file(file_path)
                out_file_part = get_file_name_without_type(file_path)

                (
                    out_file_paths,
                    out_message,
                    key_string,
                    log_files_output_paths,
                    comprehend_query_number,
                ) = tabular_anonymise_wrapper_func(
                    anon_file,
                    anon_df,
                    chosen_cols,
                    out_file_paths,
                    out_file_part,
                    out_message,
                    sheet_name,
                    anon_strategy,
                    language,
                    chosen_redact_entities,
                    in_allow_list,
                    file_type,
                    "",
                    log_files_output_paths,
                    in_deny_list,
                    max_fuzzy_spelling_mistakes_num,
                    pii_identification_method,
                    language,
                    chosen_redact_comprehend_entities,
                    comprehend_query_number,
                    comprehend_client,
                    output_folder=output_folder,
                    do_initial_clean=do_initial_clean,
                )

        # Increase latest file completed count unless we are at the last file
        if latest_file_completed != len(file_paths):
            print("Completed file number:", str(latest_file_completed))
            latest_file_completed += 1

        toc = time.perf_counter()
        out_time_float = toc - tic
        out_time = f"in {out_time_float:0.1f} seconds."
        print(out_time)

        actual_time_taken_number += out_time_float

        if isinstance(out_message, str):
            out_message = [out_message]

        out_message.append(
            "Anonymisation of file '" + out_file_part + "' successfully completed in"
        )

        out_message_out = "\n".join(out_message)
        out_message_out = out_message_out + " " + out_time

        if anon_strategy == "encrypt":
            out_message_out.append(". Your decryption key is " + key_string)

        out_message_out = (
            out_message_out
            + "\n\nPlease give feedback on the results below to help improve this app."
        )

        from tools.secure_regex_utils import safe_remove_leading_newlines

        out_message_out = safe_remove_leading_newlines(out_message_out)
        out_message_out = out_message_out.lstrip(". ")

    return (
        out_message_out,
        out_file_paths,
        out_file_paths,
        latest_file_completed,
        log_files_output_paths,
        log_files_output_paths,
        actual_time_taken_number,
        comprehend_query_number,
    )


def tabular_anonymise_wrapper_func(
    anon_file: str,
    anon_df: pd.DataFrame,
    chosen_cols: List[str],
    out_file_paths: List[str],
    out_file_part: str,
    out_message: str,
    excel_sheet_name: str,
    anon_strategy: str,
    language: str,
    chosen_redact_entities: List[str],
    in_allow_list: List[str],
    file_type: str,
    anon_xlsx_export_file_name: str,
    log_files_output_paths: List[str],
    in_deny_list: List[str] = list(),
    max_fuzzy_spelling_mistakes_num: int = 0,
    pii_identification_method: str = "Local",
    comprehend_language: Optional[str] = None,
    chosen_redact_comprehend_entities: List[str] = list(),
    comprehend_query_number: int = 0,
    comprehend_client: botocore.client.BaseClient = "",
    nlp_analyser: AnalyzerEngine = nlp_analyser,
    output_folder: str = OUTPUT_FOLDER,
    do_initial_clean: bool = DO_INITIAL_TABULAR_DATA_CLEAN,
):
    """
    This function wraps the anonymisation process for a given dataframe. It filters the dataframe based on chosen columns, applies the specified anonymisation strategy using the anonymise_script function, and exports the anonymised data to a file.

    Input Variables:
    - anon_file: The path to the file containing the data to be anonymized.
    - anon_df: The pandas DataFrame containing the data to be anonymized.
    - chosen_cols: A list of column names to be anonymized.
    - out_file_paths: A list of paths where the anonymized files will be saved.
    - out_file_part: A part of the output file name.
    - out_message: A message to be displayed during the anonymization process.
    - excel_sheet_name: The name of the Excel sheet where the anonymized data will be exported.
    - anon_strategy: The anonymization strategy to be applied.
    - language: The language of the data to be anonymized.
    - chosen_redact_entities: A list of entities to be redacted.
    - in_allow_list: A list of allowed values.
    - file_type: The type of file to be exported.
    - anon_xlsx_export_file_name: The name of the anonymized Excel file.
    - log_files_output_paths: A list of paths where the log files will be saved.
    - in_deny_list: List of specific terms to remove from the data.
    - max_fuzzy_spelling_mistakes_num (int, optional): The maximum number of spelling mistakes allowed in a searched phrase for fuzzy matching. Can range from 0-9.
    - pii_identification_method (str, optional): The method to redact personal information. Either 'Local' (spacy model), or 'AWS Comprehend' (AWS Comprehend API).
    - chosen_redact_comprehend_entities (List[str]): A list of entity types to redact from files, chosen from the official list from AWS Comprehend service.
    - comprehend_query_number (int, optional): A counter tracking the number of queries to AWS Comprehend.
    - comprehend_client (optional): The client object from AWS containing a client connection to AWS Comprehend if that option is chosen on the first tab.
    - output_folder: The folder where the anonymized files will be saved. Defaults to the 'output_folder' variable.
    - do_initial_clean (bool, optional): Whether to perform an initial cleaning of the text. Defaults to True.
    """

    def check_lists(list1, list2):
        return any(string in list2 for string in list1)

    def get_common_strings(list1, list2):
        """
        Finds the common strings between two lists.

        Args:
            list1: The first list of strings.
            list2: The second list of strings.

        Returns:
            A list containing the common strings.
        """
        common_strings = list()
        for string in list1:
            if string in list2:
                common_strings.append(string)
        return common_strings

    if pii_identification_method == "AWS Comprehend" and comprehend_client == "":
        raise (
            "Connection to AWS Comprehend service not found, please check connection details."
        )

    # Check for chosen col, skip file if not found
    all_cols_original_order = list(anon_df.columns)

    any_cols_found = check_lists(chosen_cols, all_cols_original_order)

    if any_cols_found is False:
        out_message = "No chosen columns found in dataframe: " + out_file_part
        key_string = ""
        print(out_message)
        return (
            out_file_paths,
            out_message,
            key_string,
            log_files_output_paths,
            comprehend_query_number,
        )
    else:
        chosen_cols_in_anon_df = get_common_strings(
            chosen_cols, all_cols_original_order
        )

    # Split dataframe to keep only selected columns
    # print("Remaining columns to redact:", chosen_cols_in_anon_df)

    if not anon_df.index.is_unique:
        anon_df = anon_df.reset_index(drop=True)

    anon_df_part = anon_df[chosen_cols_in_anon_df]
    anon_df_remain = anon_df.drop(chosen_cols_in_anon_df, axis=1)

    row_count = anon_df_part.shape[0]

    if row_count > MAX_TABLE_ROWS:
        out_message = f"Number of rows in dataframe is greater than {MAX_TABLE_ROWS}. Please submit a smaller dataframe."
        print(out_message)
        raise Exception(out_message)

    column_count = anon_df_part.shape[1]

    if column_count > MAX_TABLE_COLUMNS:
        out_message = f"Number of columns in dataframe is greater than {MAX_TABLE_COLUMNS}. Please submit a smaller dataframe."
        print(out_message)
        raise Exception(out_message)

    # Anonymise the selected columns
    (
        anon_df_part_out,
        key_string,
        decision_process_output_str,
        comprehend_query_number,
        decision_process_output_df,
    ) = anonymise_script(
        anon_df_part,
        anon_strategy,
        language,
        chosen_redact_entities,
        in_allow_list,
        in_deny_list,
        max_fuzzy_spelling_mistakes_num,
        pii_identification_method,
        chosen_redact_comprehend_entities,
        comprehend_query_number,
        comprehend_client,
        nlp_analyser=nlp_analyser,
        do_initial_clean=do_initial_clean,
    )

    anon_df_part_out.replace("^nan$", "", regex=True, inplace=True)

    # Rejoin the dataframe together
    anon_df_out = pd.concat([anon_df_part_out, anon_df_remain], axis=1)
    anon_df_out = anon_df_out[all_cols_original_order]

    # Export file
    #  Rename anonymisation strategy for file path naming
    if anon_strategy == "replace with 'REDACTED'":
        anon_strat_txt = "redact_replace"
    elif anon_strategy == "replace with <ENTITY_NAME>":
        anon_strat_txt = "redact_entity_type"
    elif anon_strategy == "redact completely":
        anon_strat_txt = "redact_remove"
    else:
        anon_strat_txt = anon_strategy

    # If the file is an xlsx, add a new sheet to the existing xlsx. Otherwise, write to csv
    if file_type == "xlsx":

        anon_export_file_name = anon_xlsx_export_file_name

        if not os.path.exists(anon_xlsx_export_file_name):
            wb = Workbook()
            ws = wb.active  # Get the default active sheet
            ws.title = excel_sheet_name
            wb.save(anon_xlsx_export_file_name)

        # Create a Pandas Excel writer using XlsxWriter as the engine.
        with pd.ExcelWriter(
            anon_xlsx_export_file_name,
            engine="openpyxl",
            mode="a",
            if_sheet_exists="replace",
        ) as writer:
            # Write each DataFrame to a different worksheet.
            anon_df_out.to_excel(writer, sheet_name=excel_sheet_name, index=None)

        decision_process_log_output_file = (
            anon_xlsx_export_file_name + "_" + excel_sheet_name + "_log.csv"
        )

        decision_process_output_df.to_csv(
            decision_process_log_output_file, index=None, encoding="utf-8-sig"
        )

    else:
        anon_export_file_name = (
            output_folder + out_file_part + "_anon_" + anon_strat_txt + ".csv"
        )
        anon_df_out.to_csv(anon_export_file_name, index=None, encoding="utf-8-sig")

        decision_process_log_output_file = anon_export_file_name + "_log.csv"

        decision_process_output_df.to_csv(
            decision_process_log_output_file, index=None, encoding="utf-8-sig"
        )

    out_file_paths.append(anon_export_file_name)
    out_file_paths.append(decision_process_log_output_file)

    # As files are created in a loop, there is a risk of duplicate file names being output. Use set to keep uniques.
    out_file_paths = list(set(out_file_paths))

    # Print result text to output text box if just anonymising open text
    if anon_file == "open_text":
        out_message = ["'" + anon_df_out["text"][0] + "'"]

    return (
        out_file_paths,
        out_message,
        key_string,
        log_files_output_paths,
        comprehend_query_number,
    )


def anonymise_script(
    df: pd.DataFrame,
    anon_strategy: str,
    language: str,
    chosen_redact_entities: List[str],
    in_allow_list: List[str] = list(),
    in_deny_list: List[str] = list(),
    max_fuzzy_spelling_mistakes_num: int = 0,
    pii_identification_method: str = "Local",
    chosen_redact_comprehend_entities: List[str] = list(),
    comprehend_query_number: int = 0,
    comprehend_client: botocore.client.BaseClient = "",
    custom_entities: List[str] = custom_entities,
    nlp_analyser: AnalyzerEngine = nlp_analyser,
    do_initial_clean: bool = DO_INITIAL_TABULAR_DATA_CLEAN,
    progress: Progress = Progress(track_tqdm=True),
):
    """
    Conduct anonymisation of a dataframe using Presidio and/or AWS Comprehend if chosen.

    Args:
        df (pd.DataFrame): The input DataFrame containing text to be anonymised.
        anon_strategy (str): The anonymisation strategy to apply (e.g., "replace with 'REDACTED'", "replace with <ENTITY_NAME>", "redact completely").
        language (str): The language of the text for analysis (e.g., "en", "es").
        chosen_redact_entities (List[str]): A list of entity types to redact using the local (Presidio) method.
        in_allow_list (List[str], optional): A list of terms to explicitly allow and not redact. Defaults to an empty list.
        in_deny_list (List[str], optional): A list of terms to explicitly deny and always redact. Defaults to an empty list.
        max_fuzzy_spelling_mistakes_num (int, optional): The maximum number of fuzzy spelling mistakes to tolerate for custom recognizers. Defaults to 0.
        pii_identification_method (str, optional): The method for PII identification ("Local", "AWS Comprehend", or "Both"). Defaults to "Local".
        chosen_redact_comprehend_entities (List[str], optional): A list of entity types to redact using AWS Comprehend. Defaults to an empty list.
        comprehend_query_number (int, optional): The number of queries to send to AWS Comprehend per batch. Defaults to 0.
        comprehend_client (botocore.client.BaseClient, optional): An initialized AWS Comprehend client. Defaults to an empty string.
        custom_entities (List[str], optional): A list of custom entities to be recognized. Defaults to `custom_entities`.
        nlp_analyser (AnalyzerEngine, optional): The Presidio AnalyzerEngine instance to use. Defaults to `nlp_analyser`.
        do_initial_clean (bool, optional): Whether to perform an initial cleaning of the text. Defaults to True.
        progress (Progress, optional): Gradio Progress object for tracking progress. Defaults to Progress(track_tqdm=False).
    """

    print("Identifying personal information")
    analyse_tic = time.perf_counter()

    # Initialize analyzer_results as an empty dictionary to store results by column
    results_by_column = dict()
    key_string = ""

    if isinstance(in_allow_list, list):
        if in_allow_list:
            in_allow_list_flat = in_allow_list
        else:
            in_allow_list_flat = list()
    elif isinstance(in_allow_list, pd.DataFrame):
        if not in_allow_list.empty:
            in_allow_list_flat = list(in_allow_list.iloc[:, 0].unique())
        else:
            in_allow_list_flat = list()
    else:
        in_allow_list_flat = list()

    ### Language check - check if selected language packs exist
    try:
        if language != "en":
            progress(0.1, desc=f"Loading spaCy model for {language}")

        load_spacy_model(language)

    except Exception as e:
        out_message = f"Error downloading language packs for {language}: {e}"
        print(out_message)
        raise Exception(out_message)

    # Try updating the supported languages for the spacy analyser
    try:
        nlp_analyser = create_nlp_analyser(language, existing_nlp_analyser=nlp_analyser)
        # Check list of nlp_analyser recognisers and languages
        if language != "en":
            gr.Info(
                f"Language: {language} only supports the following entity detection: {str(nlp_analyser.registry.get_supported_entities(languages=[language]))}"
            )

    except Exception as e:
        out_message = f"Error creating nlp_analyser for {language}: {e}"
        print(out_message)
        raise Exception(out_message)

    if isinstance(in_deny_list, pd.DataFrame):
        if not in_deny_list.empty:
            in_deny_list = in_deny_list.iloc[:, 0].tolist()
        else:
            # Handle the case where the DataFrame is empty
            in_deny_list = list()  # or some default value

        # Sort the strings in order from the longest string to the shortest
        in_deny_list = sorted(in_deny_list, key=len, reverse=True)

    if in_deny_list:
        nlp_analyser.registry.remove_recognizer("CUSTOM")
        new_custom_recogniser = custom_word_list_recogniser(in_deny_list)
        nlp_analyser.registry.add_recognizer(new_custom_recogniser)

        nlp_analyser.registry.remove_recognizer("CustomWordFuzzyRecognizer")
        new_custom_fuzzy_recogniser = CustomWordFuzzyRecognizer(
            supported_entities=["CUSTOM_FUZZY"],
            custom_list=in_deny_list,
            spelling_mistakes_max=in_deny_list,
            search_whole_phrase=max_fuzzy_spelling_mistakes_num,
        )
        nlp_analyser.registry.add_recognizer(new_custom_fuzzy_recogniser)

    # analyzer = nlp_analyser #AnalyzerEngine()
    batch_analyzer = BatchAnalyzerEngine(analyzer_engine=nlp_analyser)
    anonymizer = (
        AnonymizerEngine()
    )  # conflict_resolution=ConflictResolutionStrategy.MERGE_SIMILAR_OR_CONTAINED)
    batch_anonymizer = BatchAnonymizerEngine(anonymizer_engine=anonymizer)
    analyzer_results = list()

    if do_initial_clean:
        progress(0.2, desc="Cleaning text")
        for col in progress.tqdm(df.columns, desc="Cleaning text", unit="Columns"):
            df[col] = initial_clean(df[col])

    # DataFrame to dict
    df_dict = df.to_dict(orient="list")

    if pii_identification_method == "Local":

        # Use custom analyzer to be able to track progress with Gradio
        custom_results = analyze_dict(
            batch_analyzer,
            df_dict,
            language=language,
            entities=chosen_redact_entities,
            score_threshold=score_threshold,
            return_decision_process=True,
            allow_list=in_allow_list_flat,
        )

        # Initialize results_by_column with custom entity results
        for result in custom_results:
            results_by_column[result.key] = result

        # Convert the dictionary of results back to a list
        analyzer_results = list(results_by_column.values())

    # AWS Comprehend calls
    elif pii_identification_method == "AWS Comprehend" and comprehend_client:

        # Only run Local anonymisation for entities that are not covered by AWS Comprehend
        if custom_entities:
            custom_redact_entities = [
                entity
                for entity in chosen_redact_comprehend_entities
                if entity in custom_entities
            ]
            if custom_redact_entities:
                # Get results from analyze_dict
                custom_results = analyze_dict(
                    batch_analyzer,
                    df_dict,
                    language=language,
                    entities=custom_redact_entities,
                    score_threshold=score_threshold,
                    return_decision_process=True,
                    allow_list=in_allow_list_flat,
                )

                # Initialize results_by_column with custom entity results
                for result in custom_results:
                    results_by_column[result.key] = result

        max_retries = 3
        retry_delay = 3

        # Process each text column in the dictionary
        for column_name, texts in progress.tqdm(
            df_dict.items(), desc="Querying AWS Comprehend service.", unit="Columns"
        ):
            # Get or create DictAnalyzerResult for this column
            if column_name in results_by_column:
                column_results = results_by_column[column_name]
            else:
                column_results = DictAnalyzerResult(
                    recognizer_results=[[] for _ in texts], key=column_name, value=texts
                )

            # Process each text in the column
            for text_idx, text in progress.tqdm(
                enumerate(texts), desc="Querying AWS Comprehend service.", unit="Row"
            ):

                for attempt in range(max_retries):
                    try:
                        response = comprehend_client.detect_pii_entities(
                            Text=str(text), LanguageCode=language
                        )

                        comprehend_query_number += 1

                        # Add all entities from this text to the column's recognizer_results
                        for entity in response["Entities"]:
                            if (
                                entity.get("Type")
                                not in chosen_redact_comprehend_entities
                            ):
                                continue

                            recognizer_result = RecognizerResult(
                                entity_type=entity["Type"],
                                start=entity["BeginOffset"],
                                end=entity["EndOffset"],
                                score=entity["Score"],
                            )
                            column_results.recognizer_results[text_idx].append(
                                recognizer_result
                            )

                        break  # Success, exit retry loop

                    except Exception as e:
                        if attempt == max_retries - 1:
                            print(
                                f"AWS Comprehend calls failed for text: {text[:100]}... due to",
                                e,
                            )
                            raise
                        time.sleep(retry_delay)

            # Store or update the column results
            results_by_column[column_name] = column_results

        # Convert the dictionary of results back to a list
        analyzer_results = list(results_by_column.values())

    elif (pii_identification_method == "AWS Comprehend") & (not comprehend_client):
        raise ("Unable to redact, Comprehend connection details not found.")

    else:
        print("Unable to redact.")

    # Usage in the main function:
    decision_process_output_str, decision_process_output_df = generate_log(
        analyzer_results, df_dict
    )

    analyse_toc = time.perf_counter()
    analyse_time_out = (
        f"Analysing the text took {analyse_toc - analyse_tic:0.1f} seconds."
    )
    print(analyse_time_out)

    # Set up the anonymization configuration WITHOUT DATE_TIME
    simple_replace_config = {
        "DEFAULT": OperatorConfig("replace", {"new_value": "REDACTED"})
    }
    replace_config = {"DEFAULT": OperatorConfig("replace")}
    redact_config = {"DEFAULT": OperatorConfig("redact")}
    hash_config = {"DEFAULT": OperatorConfig("hash")}
    mask_config = {
        "DEFAULT": OperatorConfig(
            "mask", {"masking_char": "*", "chars_to_mask": 100, "from_end": True}
        )
    }
    people_encrypt_config = {
        "PERSON": OperatorConfig("encrypt", {"key": key_string})
    }  # The encryption is using AES cypher in CBC mode and requires a cryptographic key as an input for both the encryption and the decryption.
    fake_first_name_config = {
        "PERSON": OperatorConfig("custom", {"lambda": fake_first_name})
    }

    if anon_strategy == "replace with 'REDACTED'":
        chosen_mask_config = simple_replace_config
    elif anon_strategy == "replace_redacted":
        chosen_mask_config = simple_replace_config
    elif anon_strategy == "replace with <ENTITY_NAME>":
        chosen_mask_config = replace_config
    elif anon_strategy == "entity_type":
        chosen_mask_config = replace_config
    elif anon_strategy == "redact completely":
        chosen_mask_config = redact_config
    elif anon_strategy == "redact":
        chosen_mask_config = redact_config
    elif anon_strategy == "hash":
        chosen_mask_config = hash_config
    elif anon_strategy == "mask":
        chosen_mask_config = mask_config
    elif anon_strategy == "encrypt":
        chosen_mask_config = people_encrypt_config
        key = secrets.token_bytes(16)  # 128 bits = 16 bytes
        key_string = base64.b64encode(key).decode("utf-8")

        # Now inject the key into the operator config
        for entity, operator in chosen_mask_config.items():
            if operator.operator_name == "encrypt":
                operator.params = {"key": key_string}
    elif anon_strategy == "fake_first_name":
        chosen_mask_config = fake_first_name_config
    else:
        print("Anonymisation strategy not found. Redacting completely by default.")
        chosen_mask_config = redact_config  # Redact completely by default

    combined_config = {**chosen_mask_config}

    anonymizer_results = batch_anonymizer.anonymize_dict(
        analyzer_results, operators=combined_config
    )

    scrubbed_df = pd.DataFrame(anonymizer_results)

    return (
        scrubbed_df,
        key_string,
        decision_process_output_str,
        comprehend_query_number,
        decision_process_output_df,
    )
