"""DOCX first-page preview helpers."""

import docx
import pytest

pytest.importorskip("regex")

from tools.data_anonymise import (
    _accumulate_text_blocks_up_to_chars,
    _docx_first_page_redacted_preview,
    _extract_docx_text_blocks_with_pages,
)


def test_docx_first_page_redacted_preview_joins_page_one_blocks():
    pages = [1, 1, 2, 2]
    texts = ["Header", "Body paragraph", "Page two", "More"]
    preview = _docx_first_page_redacted_preview(pages, texts)
    assert preview == "Header\n\n\nBody paragraph"


def test_docx_first_page_redacted_preview_caps_when_no_page_breaks():
    pages = [1] * 20
    texts = [f"Block {i} " + ("x" * 200) for i in range(20)]
    preview = _docx_first_page_redacted_preview(pages, texts)
    assert "Block 0" in preview
    assert "Block 19" not in preview
    assert len(preview) <= 2600


def test_accumulate_text_blocks_up_to_chars():
    texts = ["a" * 1000, "b" * 1000, "c" * 1000]
    selected = _accumulate_text_blocks_up_to_chars(texts, 1500)
    assert len(selected) == 1
    assert selected[0] == "a" * 1000


def test_extract_docx_text_blocks_deduplicates_merged_table_cells():
    doc = docx.Document()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=1, cols=2)
    merged = table.cell(0, 0)
    merged.merge(table.cell(0, 1))
    merged.text = "Merged header"
    table.add_row()
    table.cell(1, 0).text = "Left"
    table.cell(1, 1).text = "Right"

    blocks = _extract_docx_text_blocks_with_pages(doc)
    texts = [text for _, text, _ in blocks]

    assert texts.count("Merged header") == 1
    assert texts == ["Intro", "Merged header", "Left", "Right"]
