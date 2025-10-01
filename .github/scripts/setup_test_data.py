#!/usr/bin/env python3
"""
Setup script for GitHub Actions test data.
Creates dummy test files when example data is not available.
"""

import os
import sys

import pandas as pd


def create_directories():
    """Create necessary directories."""
    dirs = ["example_data", "example_data/example_outputs"]

    for dir_path in dirs:
        os.makedirs(dir_path, exist_ok=True)
        print(f"Created directory: {dir_path}")


def create_dummy_pdf():
    """Create dummy PDFs for testing."""

    # Install reportlab if not available
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        import subprocess

        subprocess.check_call(["pip", "install", "reportlab"])
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

    try:
        # Create the main test PDF
        pdf_path = (
            "example_data/example_of_emails_sent_to_a_professor_before_applying.pdf"
        )
        print(f"Creating PDF: {pdf_path}")
        print(f"Directory exists: {os.path.exists('example_data')}")

        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "This is a test document for redaction testing.")
        c.drawString(100, 700, "Email: test@example.com")
        c.drawString(100, 650, "Phone: 123-456-7890")
        c.drawString(100, 600, "Name: John Doe")
        c.drawString(100, 550, "Address: 123 Test Street, Test City, TC 12345")
        c.showPage()

        # Add second page
        c.drawString(100, 750, "Second page content")
        c.drawString(100, 700, "More test data: jane.doe@example.com")
        c.drawString(100, 650, "Another phone: 987-654-3210")
        c.save()

        print(f"Created dummy PDF: {pdf_path}")

        # Create Partnership Agreement Toolkit PDF
        partnership_pdf_path = "example_data/Partnership-Agreement-Toolkit_0_0.pdf"
        print(f"Creating PDF: {partnership_pdf_path}")
        c = canvas.Canvas(partnership_pdf_path, pagesize=letter)
        c.drawString(100, 750, "Partnership Agreement Toolkit")
        c.drawString(100, 700, "This is a test partnership agreement document.")
        c.drawString(100, 650, "Contact: partnership@example.com")
        c.drawString(100, 600, "Phone: (555) 123-4567")
        c.drawString(100, 550, "Address: 123 Partnership Street, City, State 12345")
        c.showPage()

        # Add second page
        c.drawString(100, 750, "Page 2 - Partnership Details")
        c.drawString(100, 700, "More partnership information here.")
        c.drawString(100, 650, "Contact: info@partnership.org")
        c.showPage()

        # Add third page
        c.drawString(100, 750, "Page 3 - Terms and Conditions")
        c.drawString(100, 700, "Terms and conditions content.")
        c.drawString(100, 650, "Legal contact: legal@partnership.org")
        c.save()

        print(f"Created dummy PDF: {partnership_pdf_path}")

        # Create Graduate Job Cover Letter PDF
        cover_letter_pdf_path = "example_data/graduate-job-example-cover-letter.pdf"
        print(f"Creating PDF: {cover_letter_pdf_path}")
        c = canvas.Canvas(cover_letter_pdf_path, pagesize=letter)
        c.drawString(100, 750, "Cover Letter Example")
        c.drawString(100, 700, "Dear Hiring Manager,")
        c.drawString(100, 650, "I am writing to apply for the position.")
        c.drawString(100, 600, "Contact: applicant@example.com")
        c.drawString(100, 550, "Phone: (555) 987-6543")
        c.drawString(100, 500, "Address: 456 Job Street, Employment City, EC 54321")
        c.drawString(100, 450, "Sincerely,")
        c.drawString(100, 400, "John Applicant")
        c.save()

        print(f"Created dummy PDF: {cover_letter_pdf_path}")

    except ImportError:
        print("ReportLab not available, skipping PDF creation")
        # Create simple text files instead
        with open(
            "example_data/example_of_emails_sent_to_a_professor_before_applying.pdf",
            "w",
        ) as f:
            f.write("This is a dummy PDF file for testing")

        with open(
            "example_data/Partnership-Agreement-Toolkit_0_0.pdf",
            "w",
        ) as f:
            f.write("This is a dummy Partnership Agreement PDF file for testing")

        with open(
            "example_data/graduate-job-example-cover-letter.pdf",
            "w",
        ) as f:
            f.write("This is a dummy cover letter PDF file for testing")

        print("Created dummy text files instead of PDFs")


def create_dummy_csv():
    """Create dummy CSV files for testing."""
    # Main CSV
    csv_data = {
        "Case Note": [
            "Client visited for consultation regarding housing issues",
            "Follow-up appointment scheduled for next week",
            "Documentation submitted for review",
        ],
        "Client": ["John Smith", "Jane Doe", "Bob Johnson"],
        "Date": ["2024-01-15", "2024-01-16", "2024-01-17"],
    }
    df = pd.DataFrame(csv_data)
    df.to_csv("example_data/combined_case_notes.csv", index=False)
    print("Created dummy CSV: example_data/combined_case_notes.csv")

    # Lambeth CSV
    lambeth_data = {
        "text": [
            "Lambeth 2030 vision document content",
            "Our Future Our Lambeth strategic plan",
            "Community engagement and development",
        ],
        "page": [1, 2, 3],
    }
    df_lambeth = pd.DataFrame(lambeth_data)
    df_lambeth.to_csv(
        "example_data/Lambeth_2030-Our_Future_Our_Lambeth.pdf.csv", index=False
    )
    print("Created dummy CSV: example_data/Lambeth_2030-Our_Future_Our_Lambeth.pdf.csv")


def create_dummy_word_doc():
    """Create dummy Word document."""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document for Redaction", 0)
        doc.add_paragraph("This is a test document for redaction testing.")
        doc.add_paragraph("Contact Information:")
        doc.add_paragraph("Email: test@example.com")
        doc.add_paragraph("Phone: 123-456-7890")
        doc.add_paragraph("Name: John Doe")
        doc.add_paragraph("Address: 123 Test Street, Test City, TC 12345")

        doc.save("example_data/Bold minimalist professional cover letter.docx")
        print("Created dummy Word document")

    except ImportError:
        print("python-docx not available, skipping Word document creation")


def create_allow_deny_lists():
    """Create dummy allow/deny lists."""
    # Allow lists
    allow_data = {"word": ["test", "example", "document"]}
    pd.DataFrame(allow_data).to_csv(
        "example_data/test_allow_list_graduate.csv", index=False
    )
    pd.DataFrame(allow_data).to_csv(
        "example_data/test_allow_list_partnership.csv", index=False
    )
    print("Created allow lists")

    # Deny lists
    deny_data = {"word": ["sensitive", "confidential", "private"]}
    pd.DataFrame(deny_data).to_csv(
        "example_data/partnership_toolkit_redact_custom_deny_list.csv", index=False
    )
    pd.DataFrame(deny_data).to_csv(
        "example_data/Partnership-Agreement-Toolkit_test_deny_list_para_single_spell.csv",
        index=False,
    )
    print("Created deny lists")

    # Whole page redaction list
    page_data = {"page": [1, 2]}
    pd.DataFrame(page_data).to_csv(
        "example_data/partnership_toolkit_redact_some_pages.csv", index=False
    )
    print("Created whole page redaction list")


def create_ocr_output():
    """Create dummy OCR output CSV."""
    ocr_data = {
        "page": [1, 2, 3],
        "text": [
            "This is page 1 content with some text",
            "This is page 2 content with different text",
            "This is page 3 content with more text",
        ],
        "left": [0.1, 0.3, 0.5],
        "top": [0.95, 0.92, 0.88],
        "width": [0.05, 0.02, 0.02],
        "height": [0.01, 0.02, 0.02],
        "line": [1, 2, 3],
    }
    df = pd.DataFrame(ocr_data)
    df.to_csv(
        "example_data/example_outputs/doubled_output_joined.pdf_ocr_output.csv",
        index=False,
    )
    print("Created dummy OCR output CSV")


def create_dummy_image():
    """Create dummy image for testing."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new("RGB", (800, 600), color="white")
        draw = ImageDraw.Draw(img)

        # Try to use a system font
        try:
            font = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20
            )
        except Exception as e:
            print(f"Error loading DejaVuSans font: {e}")
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 20)
            except Exception as e:
                print(f"Error loading Arial font: {e}")
                font = ImageFont.load_default()

        # Add text to image
        draw.text((50, 50), "Test Document for Redaction", fill="black", font=font)
        draw.text((50, 100), "Email: test@example.com", fill="black", font=font)
        draw.text((50, 150), "Phone: 123-456-7890", fill="black", font=font)
        draw.text((50, 200), "Name: John Doe", fill="black", font=font)
        draw.text((50, 250), "Address: 123 Test Street", fill="black", font=font)

        img.save("example_data/example_complaint_letter.jpg")
        print("Created dummy image")

    except ImportError:
        print("PIL not available, skipping image creation")


def main():
    """Main setup function."""
    print("Setting up test data for GitHub Actions...")
    print(f"Current working directory: {os.getcwd()}")
    print(f"Python version: {sys.version}")

    create_directories()
    create_dummy_pdf()
    create_dummy_csv()
    create_dummy_word_doc()
    create_allow_deny_lists()
    create_ocr_output()
    create_dummy_image()

    print("\nTest data setup complete!")
    print("Created files:")
    for root, dirs, files in os.walk("example_data"):
        for file in files:
            file_path = os.path.join(root, file)
            print(f"  {file_path}")
            # Verify the file exists and has content
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                print(f"    Size: {file_size} bytes")
            else:
                print("    WARNING: File does not exist!")

    # Verify critical files exist
    critical_files = [
        "example_data/Partnership-Agreement-Toolkit_0_0.pdf",
        "example_data/graduate-job-example-cover-letter.pdf",
        "example_data/example_of_emails_sent_to_a_professor_before_applying.pdf",
    ]

    print("\nVerifying critical test files:")
    for file_path in critical_files:
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"✅ {file_path} exists ({file_size} bytes)")
        else:
            print(f"❌ {file_path} MISSING!")


if __name__ == "__main__":
    main()
